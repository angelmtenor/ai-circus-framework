"""Best-effort text extraction from a session-only chat attachment.

Called by `POST /documents/extract` (see `platform_registry.api`) — the file is never
written to SeaweedFS/disk beyond FastAPI's own transient `UploadFile` spooling; the caller
(ui-react's ChatPanel) discards it once this response comes back. Because of that
"good enough for one chat turn" scope, PDF OCR uses `pytesseract`/`pdf2image` (small
apt packages, no ML weights) rather than a heavier layout-aware OCR stack like
`docling` — see the root plan this was built from for that trade-off.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Literal

import pytesseract
from docx import Document
from pdf2image import convert_from_bytes
from pypdf import PdfReader

# Past this many extracted characters, cut off and flag `truncated=True` rather than
# send an unbounded blob into a chat prompt (which would blow past most models'
# context window on its own before the user's actual question is even considered).
MAX_EXTRACTED_CHARS = 50_000

# A PDF page's text layer shorter than this (after stripping whitespace) is treated as
# "no real text layer" — e.g. a scanned page with only a stray header/footer OCR'd by
# the PDF's own metadata — and falls back to image OCR for that page.
_MIN_PAGE_TEXT_CHARS = 20

DocumentKind = Literal["text", "markdown", "docx", "pdf", "image"]

_EXTENSION_KINDS: dict[str, DocumentKind] = {
    "txt": "text",
    "md": "markdown",
    "docx": "docx",
    "pdf": "pdf",
    "png": "image",
    "jpg": "image",
    "jpeg": "image",
    "webp": "image",
}


class UnsupportedDocumentError(ValueError):
    """Raised when the uploaded file's extension isn't one this module handles."""


@dataclass(frozen=True)
class ExtractedDocument:
    """Result of extracting text from one uploaded file."""

    kind: DocumentKind
    text: str
    truncated: bool
    page_count: int | None = None
    used_ocr: bool = False


def kind_for_filename(filename: str) -> DocumentKind:
    """Classify `filename` by extension, raising `UnsupportedDocumentError` otherwise."""
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    kind = _EXTENSION_KINDS.get(extension)
    if kind is None:
        raise UnsupportedDocumentError(f"Unsupported file extension: {filename!r}")
    return kind


def _cap(text: str) -> tuple[str, bool]:
    stripped = text.strip()
    if len(stripped) <= MAX_EXTRACTED_CHARS:
        return stripped, False
    return stripped[:MAX_EXTRACTED_CHARS], True


def _extract_text_file(data: bytes) -> ExtractedDocument:
    text, truncated = _cap(data.decode("utf-8", errors="replace"))
    return ExtractedDocument(kind="text", text=text, truncated=truncated)


def _extract_markdown_file(data: bytes) -> ExtractedDocument:
    text, truncated = _cap(data.decode("utf-8", errors="replace"))
    return ExtractedDocument(kind="markdown", text=text, truncated=truncated)


def _extract_docx(data: bytes) -> ExtractedDocument:
    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text, truncated = _cap("\n".join(parts))
    return ExtractedDocument(kind="docx", text=text, truncated=truncated)


def _ocr_image_bytes(data: bytes) -> str:
    from PIL import Image

    # pytesseract.image_to_string's return type is overloaded on `output_type`
    # (defaults to plain str) — str() pins the type for the type checker without
    # changing behavior.
    return str(pytesseract.image_to_string(Image.open(io.BytesIO(data))))


def _extract_image(data: bytes) -> ExtractedDocument:
    text, truncated = _cap(_ocr_image_bytes(data))
    return ExtractedDocument(kind="image", text=text, truncated=truncated, used_ocr=True)


def _extract_pdf(data: bytes) -> ExtractedDocument:
    reader = PdfReader(io.BytesIO(data))
    page_count = len(reader.pages)
    page_texts: list[str] = []
    used_ocr = False
    ocr_pages: list[int] = []
    for index, page in enumerate(reader.pages):
        page_text = (page.extract_text() or "").strip()
        if len(page_text) < _MIN_PAGE_TEXT_CHARS:
            ocr_pages.append(index)
            page_texts.append("")  # placeholder, filled in below if OCR succeeds
        else:
            page_texts.append(page_text)

    if ocr_pages:
        # Rendering every page as an image up front would be wasteful for a mostly
        # text-native PDF — only render the specific pages whose text layer was empty.
        images = convert_from_bytes(data, first_page=1, last_page=page_count)
        for index in ocr_pages:
            page_texts[index] = str(pytesseract.image_to_string(images[index]))
            used_ocr = True

    text, truncated = _cap("\n\n".join(t for t in page_texts if t))
    return ExtractedDocument(kind="pdf", text=text, truncated=truncated, page_count=page_count, used_ocr=used_ocr)


def extract_document(filename: str, data: bytes) -> ExtractedDocument:
    """Dispatch `data` (the raw bytes of `filename`) to the right extractor.

    Raises `UnsupportedDocumentError` for an unrecognized extension — the caller
    turns that into an HTTP 415.
    """
    kind = kind_for_filename(filename)
    if kind == "text":
        return _extract_text_file(data)
    if kind == "markdown":
        return _extract_markdown_file(data)
    if kind == "docx":
        return _extract_docx(data)
    if kind == "image":
        return _extract_image(data)
    return _extract_pdf(data)
